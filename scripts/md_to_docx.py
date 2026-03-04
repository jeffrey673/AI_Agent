"""Convert SKIN1004_Security_Architecture.md to Word (.docx) document."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def create_security_doc():
    md_path = Path(__file__).parent.parent / "docs" / "SKIN1004_Security_Architecture.md"
    out_path = Path(__file__).parent.parent / "docs" / "SKIN1004_Security_Architecture_v2.docx"

    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "맑은 고딕"
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if level == 1:
            hs.font.size = Pt(18)
        elif level == 2:
            hs.font.size = Pt(14)
        else:
            hs.font.size = Pt(12)

    # Parse state
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    in_blockquote = False
    bq_lines = []
    i = 0

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        # Filter out separator rows
        data_rows = []
        for row in table_rows:
            cells = [c.strip() for c in row.strip("|").split("|")]
            # Skip separator rows like |---|---|
            if all(set(c.strip()) <= {"-", ":", " "} for c in cells):
                continue
            data_rows.append(cells)
        if not data_rows:
            in_table = False
            table_rows = []
            return

        ncols = max(len(r) for r in data_rows)
        tbl = doc.add_table(rows=len(data_rows), cols=ncols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row_cells in enumerate(data_rows):
            for ci in range(ncols):
                cell_text = row_cells[ci].strip() if ci < len(row_cells) else ""
                # Clean markdown formatting
                cell_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
                cell_text = re.sub(r"`(.*?)`", r"\1", cell_text)
                cell = tbl.cell(ri, ci)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles["Normal"]
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                # Header row styling
                if ri == 0:
                    shading = cell._element.get_or_add_tcPr()
                    bg = shading.makeelement(qn("w:shd"), {
                        qn("w:fill"): "1A1A2E",
                        qn("w:val"): "clear",
                    })
                    shading.append(bg)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True

        doc.add_paragraph()  # spacing
        in_table = False
        table_rows = []

    def flush_blockquote():
        nonlocal bq_lines, in_blockquote
        if not bq_lines:
            return
        text = " ".join(bq_lines)
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        in_blockquote = False
        bq_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block - add as formatted text
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                # Add light gray background via shading
                shading_elm = run._element.get_or_add_rPr()
                bg = shading_elm.makeelement(qn("w:shd"), {
                    qn("w:fill"): "F0F0F0",
                    qn("w:val"): "clear",
                })
                shading_elm.append(bg)
                in_code_block = False
                code_lines = []
            else:
                # Start code block - flush any pending content
                if in_table:
                    flush_table()
                if in_blockquote:
                    flush_blockquote()
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            if in_table:
                flush_table()
            if in_blockquote:
                flush_blockquote()
            # Add thin horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # Empty line
        if not line.strip():
            if in_table:
                flush_table()
            if in_blockquote:
                flush_blockquote()
            i += 1
            continue

        # Table row
        if line.strip().startswith("|"):
            if in_blockquote:
                flush_blockquote()
            in_table = True
            table_rows.append(line)
            i += 1
            continue

        # Flush table if we're past it
        if in_table and not line.strip().startswith("|"):
            flush_table()

        # Blockquote
        if line.strip().startswith(">"):
            text = line.strip()[1:].strip()
            if text.startswith(">"):
                text = text[1:].strip()
            bq_lines.append(text)
            in_blockquote = True
            i += 1
            continue

        if in_blockquote and not line.strip().startswith(">"):
            flush_blockquote()

        # Headings
        if line.startswith("# "):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(line[5:].strip())
            run.font.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        # Regular text
        text = line.strip()
        # Process inline formatting
        p = doc.add_paragraph()
        # Split by bold markers
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                # Handle inline code
                code_parts = re.split(r"(`.*?`)", part)
                for cp in code_parts:
                    if cp.startswith("`") and cp.endswith("`"):
                        run = p.add_run(cp[1:-1])
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    else:
                        run = p.add_run(cp)

        i += 1

    # Flush remaining
    if in_table:
        flush_table()
    if in_blockquote:
        flush_blockquote()

    doc.save(str(out_path))
    print(f"Word document saved: {out_path}")
    return out_path


if __name__ == "__main__":
    create_security_doc()
